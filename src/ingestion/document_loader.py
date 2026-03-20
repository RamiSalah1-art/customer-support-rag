"""
وحدة تحميل المستندات - تدعم تنسيقات متعددة وتتبع الملفات الجديدة
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
import PyPDF2
from docx import Document
import markdown
from loguru import logger
import hashlib
import json
from datetime import datetime

class DocumentLoader:
    """تحميل المستندات من تنسيقات مختلفة مع تتبع التغييرات"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.supported_formats = config.get('ingestion', {}).get('supported_formats', ['pdf', 'docx', 'txt', 'md'])
        self.tracking_file = Path("data/processed/file_tracking.json")
        self.processed_files = self._load_tracking()
        
    def _load_tracking(self) -> Dict[str, Any]:
        """تحميل سجل الملفات التي تمت معالجتها سابقاً"""
        if self.tracking_file.exists():
            try:
                with open(self.tracking_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return {"files": {}}
        return {"files": {}}
    
    def _save_tracking(self):
        """حفظ سجل الملفات المعالجة"""
        self.tracking_file.parent.mkdir(parents=True, exist_ok=True)
        with open(self.tracking_file, 'w', encoding='utf-8') as f:
            json.dump(self.processed_files, f, ensure_ascii=False, indent=2)
    
    def _get_file_hash(self, file_path: Path) -> str:
        """حساب هاش للملف لمعرفة إذا تغير"""
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            buf = f.read(65536)
            while len(buf) > 0:
                hasher.update(buf)
                buf = f.read(65536)
        return hasher.hexdigest()
    
    def get_new_and_modified_files(self, directory_path: str) -> tuple[List[Path], List[Path]]:
        """
        تحديد الملفات الجديدة والمعدلة
        
        Returns:
            tuple: (ملفات جديدة, ملفات معدلة)
        """
        directory = Path(directory_path)
        if not directory.exists():
            logger.warning(f"المجلد {directory_path} غير موجود، سيتم إنشاؤه")
            directory.mkdir(parents=True, exist_ok=True)
            return [], []
        
        new_files = []
        modified_files = []
        current_files = set()
        
        # فحص جميع الملفات في المجلد
        for file_path in directory.glob('**/*'):
            if not file_path.is_file():
                continue
                
            extension = file_path.suffix.lower().replace('.', '')
            if extension not in self.supported_formats:
                continue
                
            current_files.add(str(file_path))
            file_hash = self._get_file_hash(file_path)
            file_stat = file_path.stat()
            modified_time = datetime.fromtimestamp(file_stat.st_mtime).isoformat()
            
            # التحقق مما إذا كان الملف موجوداً في السجل
            if str(file_path) not in self.processed_files["files"]:
                logger.info(f"📄 ملف جديد: {file_path.name}")
                new_files.append(file_path)
                self.processed_files["files"][str(file_path)] = {
                    "hash": file_hash,
                    "modified": modified_time,
                    "processed": datetime.now().isoformat()
                }
            else:
                # التحقق مما إذا كان الملف قد تغير
                old_info = self.processed_files["files"][str(file_path)]
                if old_info["hash"] != file_hash:
                    logger.info(f"🔄 ملف معدل: {file_path.name}")
                    modified_files.append(file_path)
                    self.processed_files["files"][str(file_path)].update({
                        "hash": file_hash,
                        "modified": modified_time,
                        "processed": datetime.now().isoformat()
                    })
        
        # إزالة الملفات المحذوفة من التتبع
        files_to_remove = [f for f in self.processed_files["files"] if f not in current_files]
        for f in files_to_remove:
            logger.info(f"🗑️ ملف محذوف من التتبع: {Path(f).name}")
            del self.processed_files["files"][f]
        
        # حفظ التحديثات
        self._save_tracking()
        
        return new_files, modified_files
    
    def load_document(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        تحميل مستند بناءً على امتداده
        
        Args:
            file_path: مسار الملف
            
        Returns:
            قاموس يحتوي على النص والبيانات الوصفية أو None في حالة الخطأ
        """
        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"الملف {file_path} غير موجود")
            return None
        
        try:
            extension = file_path.suffix.lower().replace('.', '')
            
            if extension == 'pdf':
                return self._load_pdf(file_path)
            elif extension == 'docx':
                return self._load_docx(file_path)
            elif extension == 'txt':
                return self._load_txt(file_path)
            elif extension == 'md':
                return self._load_markdown(file_path)
            else:
                logger.warning(f"تنسيق غير مدعوم: {extension}")
                return None
        except Exception as e:
            logger.error(f"خطأ في تحميل {file_path}: {e}")
            return None
    
    def _load_pdf(self, file_path: Path) -> Dict[str, Any]:
        """تحميل ملف PDF"""
        text = ""
        metadata = {
            'source': str(file_path),
            'format': 'pdf',
            'filename': file_path.name,
            'processed': datetime.now().isoformat()
        }
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            metadata['num_pages'] = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text += f"\n--- صفحة {page_num + 1} ---\n"
                    text += page_text
        
        return {'text': text, 'metadata': metadata}
    
    def _load_docx(self, file_path: Path) -> Dict[str, Any]:
        """تحميل ملف Word"""
        text = ""
        metadata = {
            'source': str(file_path),
            'format': 'docx',
            'filename': file_path.name,
            'processed': datetime.now().isoformat()
        }
        
        doc = Document(file_path)
        metadata['paragraph_count'] = len(doc.paragraphs)
        
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        return {'text': text, 'metadata': metadata}
    
    def _load_txt(self, file_path: Path) -> Dict[str, Any]:
        """تحميل ملف نصي"""
        metadata = {
            'source': str(file_path),
            'format': 'txt',
            'filename': file_path.name,
            'processed': datetime.now().isoformat()
        }
        
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
            metadata['char_count'] = len(text)
            metadata['line_count'] = len(text.split('\n'))
        
        return {'text': text, 'metadata': metadata}
    
    def _load_markdown(self, file_path: Path) -> Dict[str, Any]:
        """تحميل ملف Markdown"""
        metadata = {
            'source': str(file_path),
            'format': 'markdown',
            'filename': file_path.name,
            'processed': datetime.now().isoformat()
        }
        
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
            metadata['char_count'] = len(text)
        
        return {'text': text, 'metadata': metadata}
    
    def load_all_documents(self, directory_path: str, force_reload: bool = False) -> List[Dict[str, Any]]:
        """
        تحميل جميع المستندات من مجلد (جديد أو معدل فقط)
        
        Args:
            directory_path: مسار المجلد
            force_reload: إجبار إعادة تحميل كل الملفات
            
        Returns:
            قائمة بالمستندات المحملة (الجديدة فقط)
        """
        if force_reload:
            # إعادة تعيين التتبع
            self.processed_files = {"files": {}}
            self._save_tracking()
        
        new_files, modified_files = self.get_new_and_modified_files(directory_path)
        files_to_load = new_files + modified_files
        
        if not files_to_load:
            logger.info("✅ لا توجد ملفات جديدة أو معدلة")
            return []
        
        logger.info(f"📂 جاري تحميل {len(files_to_load)} ملف (جديد أو معدل)")
        
        documents = []
        for file_path in files_to_load:
            logger.info(f"📖 قراءة: {file_path.name}")
            doc = self.load_document(file_path)
            if doc:
                documents.append(doc)
        
        logger.success(f"✅ تم تحميل {len(documents)} مستند بنجاح")
        return documents